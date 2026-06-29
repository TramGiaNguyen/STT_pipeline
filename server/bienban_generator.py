"""
bienban_generator.py — Module tạo biên bản họp từ transcript

Luồng xử lý:
    1. Nhận nội dung transcript (đã parse từ .srt/.txt)
    2. Gọi Gemini 2.5 Flash phân tích → trích xuất thông tin cuộc họp
    3. Fill thông tin vào template DOCX chuẩn Viện AIDTI (mỗi phiên 1 trang riêng)
    4. Trả về bytes file .docx
"""

import copy
import json
import logging
import os
import io
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google import genai

logger = logging.getLogger(__name__)

# Đường dẫn template DOCX
TEMPLATE_PATH = Path(__file__).parent.parent / "TEMPLATE_Bien_ban_hop_chuyen_mon_Vien_AIDTI.docx"

# ==================== Prompt AI ====================

ANALYSIS_PROMPT = """Bạn là Thư ký chuyên nghiệp đang lập biên bản họp chuyên môn cho Viện Trí tuệ nhân tạo và Chuyển đổi số (AIDTI), Trường Đại học Bình Dương. Nhiệm vụ của bạn là đọc bản ghi âm cuộc họp và trích xuất thông tin để điền vào biên bản.

**YÊU CẦU NGHIÊM NGẶT VỀ VĂN PHONG VÀ ĐỊNH DẠNG:**
1. Văn phong hành chính, khách quan: Tuyệt đối KHÔNG viết theo kiểu kể chuyện hay bình luận cá nhân (VD: Không dùng "Cuộc họp diễn ra trong không khí...", "Cuộc họp được điều hành bởi..."). Sử dụng các từ ngữ chỉ đạo, giao việc như: "Giao [Tên người/đơn vị] thực hiện...", "Thống nhất...", "Tiếp tục triển khai...". Khi một việc đã có người/đơn vị nhận, PHẢI mở đầu bằng "Giao " + TÊN (KHÔNG có chữ "cho" sau "Giao"), chứ KHÔNG dùng "Yêu cầu...", "Đề nghị...".
2. Tuyệt đối KHÔNG dùng Markdown hay Bullet/Đánh số: KHÔNG sử dụng các ký tự ** hoặc * hoặc #. Tuyệt đối KHÔNG dùng gạch đầu dòng (-) hay đánh số (1., 2., a., b.).
3. XUỐNG DÒNG CHO MỖI Ý: Mỗi ý chính hoặc một công việc/chỉ đạo cụ thể PHẢI được XUỐNG DÒNG (tạo thành một dòng mới). Tuyệt đối KHÔNG gộp chung nhiều việc vào một đoạn văn dài liên tục.
4. Mục Nội dung cuộc họp ("noi_dung"): CHỈ ghi rất NGẮN GỌN các chủ đề chính được thảo luận. Mỗi chủ đề XUỐNG DÒNG. Tuyệt đối KHÔNG ghi chi tiết nhiệm vụ ở phần này.
5. Mục Kết luận cuộc họp ("ket_luan"): Đây mới là nơi ghi CHI TIẾT các chỉ đạo, phân công nhiệm vụ. Gộp các nhiệm vụ thành từng nhóm (Tên mục). Nội dung bên trong trình bày rõ ràng ai làm gì. MỖI CHỈ ĐẠO/NHIỆM VỤ PHẢI XUỐNG DÒNG.
   - BẮT BUỘC cú pháp câu giao việc (để hệ thống trích xuất công việc đọc đúng): mỗi việc có người/đơn vị nhận PHẢI theo mẫu **"Giao " + [TÊN] + [ĐỘNG TỪ] + [nội dung] + [thời hạn nếu có]**, trong đó:
       • NGAY SAU "Giao" là TÊN người/đơn vị — TUYỆT ĐỐI KHÔNG chèn "cho" (viết "Giao Hiếu ...", KHÔNG viết "Giao cho Hiếu ...").
       • [ĐỘNG TỪ] đứng NGAY SAU tên và PHẢI chọn trong danh sách (chỉ dùng đúng các từ này): thực hiện, chuẩn bị, phối hợp, xây dựng, triển khai, tổ chức, rà soát, kiểm tra, cập nhật, hoàn thiện, tiếp tục, gửi, hướng dẫn, nhắc nhở, phụ trách, điều phối, đào tạo, kết nối, liên hệ, soạn thảo, nắm rõ. Nếu việc gốc dùng động từ khác (trao đổi/dọn dẹp/học/bật/nhận thông tin...), hãy diễn đạt lại bằng một động từ trong danh sách (vd "dọn dẹp" → "thực hiện dọn dẹp", "học vài câu" → "chuẩn bị một vài câu", "bật 3 TV" → "chuẩn bị bật 3 TV", "nhận thông tin" → "nắm rõ thông tin", "trao đổi" → "phối hợp").
       • Nhiều người cùng làm: ngăn cách bằng DẤU PHẨY, KHÔNG dùng "và" (viết "Giao Tài, Khang chuẩn bị ...", KHÔNG viết "Giao Tài và Khang ...").
       • Người chỉ PHỐI HỢP (không trực tiếp làm) ghi bằng cụm "phối hợp với [Tên]" trong câu.
       • TÊN phải VIẾT GIỐNG HỆT tên trong danh sách thành viên (mục "thanh_vien"), KHÔNG kèm xưng hô ("anh","chị","thầy","cô","em").
     VD ĐÚNG: "Giao Hiếu phối hợp với Quang, Nhân chuẩn bị các mô hình robot, IoT để trưng bày."; "Giao Tài, Khang chuẩn bị bật 3 TV (1 TV giữa, 2 TV cánh) để trình chiếu sản phẩm."; "Giao Huân chuẩn bị một vài câu giao tiếp tiếng Nga cơ bản để chào hỏi đoàn."
   - Với nội dung KHÔNG gắn người/đơn vị cụ thể (quyết định/thống nhất chung, thông tin tham khảo), giữ cách diễn đạt trung tính như "Thống nhất...", "Tiếp tục triển khai..." và KHÔNG được ép thêm "Giao" khi không có ai/đơn vị nhận việc.
6. Danh sách thành viên ("thanh_vien") — quan trọng cho việc đối chiếu phân công:
   - CHỈ liệt kê những người THỰC SỰ được nhắc tên trong cuộc họp. KHÔNG thêm phần tử trống/độn để đủ chỗ — bao nhiêu người thì bấy nhiêu dòng.
   - Mỗi người chỉ xuất hiện MỘT lần với MỘT tên chuẩn duy nhất. Nếu một người bị phiên âm thành nhiều biến thể do lỗi nhận dạng giọng nói (vd "Phố"/"Phú", "Khang"/"Khen"), hãy GỘP về một tên và dùng nhất quán.
   - "ho_ten" KHÔNG kèm xưng hô ("anh", "chị", "thầy", "cô", "em") — xưng hô/chức danh để ở "chuc_vu". Ưu tiên họ tên đầy đủ nếu transcript có; nếu chỉ có tên gọi thì giữ đúng tên đó.
   - BẮT BUỘC nhất quán: mọi tên xuất hiện trong "ket_luan" (sau "Giao", sau "phối hợp với"...) phải TRÙNG KHỚP TỪNG KÝ TỰ với một "ho_ten" trong "thanh_vien". Nếu một người được giao việc mà chưa có trong "thanh_vien" thì PHẢI bổ sung họ vào danh sách.
7. Trường "so_bien_ban": CHỈ ghi CON SỐ (VD: "180"). KHÔNG ghi kèm "/BB-AIDTI" hay bất kỳ hậu tố nào. Nếu không rõ, để trống "".

**QUY TẮC VỀ PHIÊN HỌP:**
- Nếu cuộc họp CHỈ CÓ 1 PHIÊN hoặc KHÔNG chia phiên, mảng "danh_sach_phien" chỉ có 1 phần tử. Trường "phien" để trống "".
- Nếu cuộc họp CÓ NHIỀU PHIÊN (Phiên I, Phiên II, ...), PHẢI tách biệt hoàn toàn từng phiên thành từng phần tử riêng trong mảng "danh_sach_phien". Mỗi phiên có đầy đủ thời gian, địa điểm, danh sách thành viên, nội dung, kết luận riêng. Trường "phien" ghi tên phiên đó (VD: "I: 02 NGÀNH", "II: CÁC TRUNG TÂM").

**TRẢ VỀ JSON THEO ĐÚNG CẤU TRÚC SAU (CHỈ TRẢ JSON, KHÔNG GIẢI THÍCH):**

```json
{
  "so_bien_ban": "",
  "lan_hop": "",
  "thang_hop": "",
  "nam_hop": "2026",
  "danh_sach_phien": [
    {
      "phien": "",
      "thoi_gian_bat_dau": "",
      "ngay": "",
      "thang": "",
      "nam": "2026",
      "dia_diem": "",
      "thanh_vien": [
        {"stt": 1, "ho_ten": "", "chuc_vu": "", "ghi_chu": ""}
      ],
      "vang_mat": "",
      "noi_dung": "Thảo luận chủ đề 1; Thảo luận chủ đề 2.",
      "ket_luan": [
        {
          "ten_muc": "Hoạt động A",
          "noi_dung": "Giao Nguyễn Văn X thực hiện Y trước ngày 30/06/2026. Thống nhất triển khai Z."
        }
      ],
      "gio_ket_thuc": "",
      "ngay_ket_thuc": ""
    }
  ]
}
```

**BẢN GHI ÂM CUỘC HỌP:**

{transcript}
"""


# ==================== Gemini API ====================

def _parse_gemini_json(response_text: str) -> dict:
    """
    Parse JSON do Gemini trả về một cách 'khoan dung'.

    Vì sao cần: prompt yêu cầu "XUỐNG DÒNG cho mỗi ý" nên Gemini hay chèn KÝ TỰ XUỐNG
    DÒNG THẬT (\\n, \\t) vào trong các chuỗi "noi_dung"/"ket_luan". json.loads mặc định
    (strict=True) coi đó là lỗi → "Invalid control character". Ở đây:
      - Gỡ rào markdown ```json ... ``` nếu model lỡ bọc.
      - Cắt từ '{' đầu tới '}' cuối (phòng khi có chữ thừa quanh JSON).
      - Dùng strict=False để CHẤP NHẬN ký tự điều khiển thô trong chuỗi.
    Các newline này về sau lại có ích: _insert_paragraph_after() tách theo '\\n' thành
    từng dòng/đoạn đúng như mong muốn.
    """
    text = response_text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text).strip()
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        text = text[start:end + 1]
    return json.loads(text, strict=False)


def analyze_transcript_with_gemini(transcript_text: str) -> dict:
    """
    Gọi Gemini 2.5 Flash để phân tích transcript cuộc họp.
    Có retry logic khi gặp lỗi 503 (quá tải) hoặc 429 (rate limit).
    
    Args:
        transcript_text: Nội dung transcript (plain text, đã bỏ timestamp)
    
    Returns:
        dict: Dữ liệu phân tích theo cấu trúc JSON ở trên
    """
    import time
    
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY chưa được cấu hình trong file .env")
    
    client = genai.Client(api_key=api_key)
    prompt = ANALYSIS_PROMPT.replace("{transcript}", transcript_text)
    
    # Danh sách model ưu tiên (fallback nếu model chính quá tải)
    models_to_try = ["gemini-2.5-flash", "gemini-2.0-flash"]
    max_retries = 3
    retry_delay = 15  # seconds
    
    last_error = None
    
    for model_name in models_to_try:
        for attempt in range(max_retries):
            try:
                logger.info(
                    f"Gọi {model_name} phân tích transcript "
                    f"({len(transcript_text)} ký tự) — lần thử {attempt + 1}/{max_retries}"
                )
                
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                    config={
                        "temperature": 0.2,
                        "response_mime_type": "application/json",
                    },
                )
                
                response_text = response.text.strip()
                logger.info(f"Gemini ({model_name}) trả về {len(response_text)} ký tự")

                analysis = _parse_gemini_json(response_text)
                return analysis
                
            except json.JSONDecodeError as e:
                logger.error(f"Lỗi parse JSON từ Gemini: {e}")
                logger.error(f"Response text: {response_text[:500]}")
                raise ValueError(f"Gemini trả về JSON không hợp lệ: {e}")
            except Exception as e:
                last_error = e
                err_str = str(e)
                # Kiểm tra nếu là lỗi tạm thời (503, 429)
                if "503" in err_str or "429" in err_str or "UNAVAILABLE" in err_str or "RESOURCE_EXHAUSTED" in err_str:
                    logger.warning(
                        f"{model_name} lỗi tạm thời (lần {attempt + 1}): {err_str[:200]}. "
                        f"Chờ {retry_delay}s rồi thử lại..."
                    )
                    time.sleep(retry_delay)
                    continue
                else:
                    # Lỗi khác, không retry
                    raise ValueError(f"Lỗi Gemini API: {err_str[:300]}")
        
        logger.warning(f"Model {model_name} thất bại sau {max_retries} lần, thử model tiếp theo...")
    
    # Nếu tất cả model đều thất bại
    raise ValueError(
        f"Không thể gọi Gemini API sau khi thử tất cả models. "
        f"Lỗi cuối: {str(last_error)[:300]}"
    )


# ==================== DOCX Generation ====================

def fill_docx_template(analysis: dict) -> bytes:
    """
    Entry point chính: tạo file DOCX từ dữ liệu phân tích.
    Hỗ trợ tự động tách nhiều phiên thành nhiều trang.
    
    Args:
        analysis: dict dữ liệu từ Gemini
    
    Returns:
        bytes: Nội dung file .docx hoàn chỉnh
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Không tìm thấy template: {TEMPLATE_PATH}")
    
    # Lấy thông tin chung (header)
    so_bb = analysis.get("so_bien_ban", "")
    # Chỉ lấy phần số, bỏ "/BB-AIDTI" nếu AI vẫn ghi thừa
    so_bb = re.sub(r'[/\\].*', '', so_bb).strip()
    
    lan_hop = analysis.get("lan_hop", "")
    thang_hop = analysis.get("thang_hop", "")
    nam_hop = analysis.get("nam_hop", "2026")
    
    # Lấy danh sách phiên
    danh_sach_phien = analysis.get("danh_sach_phien", [])
    
    # Tương thích ngược: nếu AI trả về cấu trúc cũ (không có danh_sach_phien)
    if not danh_sach_phien:
        danh_sach_phien = [analysis]
    
    if len(danh_sach_phien) == 1:
        # Chỉ 1 phiên: fill trực tiếp vào template
        doc = _fill_single_session(
            danh_sach_phien[0],
            so_bb=so_bb,
            lan_hop=lan_hop,
            thang_hop=thang_hop,
            nam_hop=nam_hop,
        )
    else:
        # Nhiều phiên: tạo mỗi phiên 1 doc rồi ghép lại
        doc = _generate_multi_session_docx(
            danh_sach_phien,
            so_bb=so_bb,
            lan_hop=lan_hop,
            thang_hop=thang_hop,
            nam_hop=nam_hop,
        )
    
    # Xuất file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info("Đã tạo file DOCX biên bản họp thành công")
    return buffer.getvalue()


def _generate_multi_session_docx(
    danh_sach_phien: list,
    so_bb: str,
    lan_hop: str,
    thang_hop: str,
    nam_hop: str,
) -> Document:
    """
    Tạo file DOCX cho cuộc họp có nhiều phiên.
    Mỗi phiên sẽ được điền vào 1 bản copy riêng của template,
    sau đó ghép lại thành 1 Document (có ngắt trang giữa các phiên).
    """
    final_doc = None
    
    for idx, phien_data in enumerate(danh_sach_phien):
        session_doc = _fill_single_session(
            phien_data,
            so_bb=so_bb,
            lan_hop=lan_hop,
            thang_hop=thang_hop,
            nam_hop=nam_hop,
        )
        
        is_last_session = (idx == len(danh_sach_phien) - 1)
        
        if not is_last_session:
            # Xóa block chữ ký và "Nơi nhận" ở các phiên không phải cuối cùng
            _remove_signature_block(session_doc)
        
        if idx == 0:
            # Phiên đầu tiên → dùng trực tiếp làm doc gốc
            final_doc = session_doc
        else:
            # Các phiên tiếp theo → chèn page break rồi copy nội dung sang
            _append_doc_with_page_break(final_doc, session_doc)
    
    return final_doc

def _remove_signature_block(doc: Document):
    """
    Xóa block chữ ký và phần "Nơi nhận" ở cuối file (sau đoạn 'Cuộc họp kết thúc lúc').
    Dùng cho các phiên họp trung gian để chữ ký chỉ xuất hiện 1 lần ở cuối file cuối cùng.
    """
    body = doc.element.body
    end_idx = -1
    
    for i, child in enumerate(body):
        if child.tag.endswith('}p'):
            text_content = ''.join(node.text or '' for node in child.iter() if node.text)
            if "Cuộc họp kết thúc lúc" in text_content:
                end_idx = i
                break
                
    if end_idx != -1:
        elements_to_remove = []
        for i in range(end_idx + 1, len(body)):
            if not body[i].tag.endswith('sectPr'):
                elements_to_remove.append(body[i])
        for elem in elements_to_remove:
            body.remove(elem)


def _fill_single_session(
    session: dict,
    so_bb: str = "",
    lan_hop: str = "",
    thang_hop: str = "",
    nam_hop: str = "2026",
) -> Document:
    """
    Điền dữ liệu của 1 phiên họp vào 1 bản copy template DOCX.
    
    Args:
        session: dict chứa thông tin 1 phiên họp
        so_bb, lan_hop, thang_hop, nam_hop: thông tin chung từ analysis gốc
    
    Returns:
        Document: Đối tượng Document đã điền dữ liệu
    """
    doc = Document(str(TEMPLATE_PATH))
    
    ngay = session.get("ngay", "")
    thang = session.get("thang", "")
    nam = session.get("nam", "2026")
    
    # ========== 1. Fill bảng tiêu đề (Table 0) ==========
    header_table = doc.tables[0]
    
    # Cột trái: Số: .../BB-AIDTI
    left_cell = header_table.rows[0].cells[0]
    if so_bb:
        _replace_in_cell_paragraphs(left_cell, "Số:         /BB-AIDTI", f"Số: {so_bb}/BB-AIDTI")
    # Nếu so_bb trống, giữ nguyên template gốc cho người dùng tự điền
    
    # Cột phải: ngày tháng năm
    right_cell = header_table.rows[0].cells[1]
    _replace_in_cell_paragraphs(
        right_cell,
        "ngày  tháng  năm  2026",
        f"ngày {ngay} tháng {thang} năm {nam}"
    )
    
    # ========== 2. Fill các paragraph chính ==========
    paragraphs_to_insert = []  # (paragraph, text) tuples chờ insert sau
    
    for para in doc.paragraphs:
        text = para.text
        
        # Tiêu đề biên bản: lần + tháng + năm
        if "Viện Trí tuệ nhân tạo và Chuyển đổi số lần" in text:
            _replace_paragraph_text(
                para,
                f"Viện Trí tuệ nhân tạo và Chuyển đổi số lần {lan_hop} tháng {thang_hop} năm {nam_hop}"
            )
        
        # Phiên họp
        elif text.strip().startswith("PHIÊN") and text.strip().endswith(":"):
            phien = session.get("phien", "")
            if phien:
                _replace_paragraph_text(para, f"PHIÊN {phien}:")
            else:
                _replace_paragraph_text(para, "")
        
        # Thời gian
        elif "Thời gian:" in text and "giờ, ngày" in text:
            thoi_gian = session.get("thoi_gian_bat_dau", "")
            _replace_paragraph_text(
                para,
                f"Thời gian: {thoi_gian} giờ, ngày {ngay} tháng {thang} năm {nam}"
            )
        
        # Địa điểm
        elif "Địa điểm:" in text and "cơ sở chính" in text:
            dia_diem = session.get("dia_diem", "")
            _replace_paragraph_text(
                para,
                f"Địa điểm: {dia_diem}, cơ sở chính Trường Đại học Bình Dương."
            )
        
        # Nhân sự vắng
        elif "Các nhân sự vắng có lý do" in text:
            vang_mat = session.get("vang_mat", "Không")
            _replace_paragraph_text(
                para,
                f"Các nhân sự vắng có lý do (trùng lịch dạy): {vang_mat}"
            )
        
        # Nội dung cuộc họp (phần ghi chú hướng dẫn)
        elif "(Ghi tóm tắt các nội dung được báo cáo" in text:
            noi_dung = session.get("noi_dung", "")
            _replace_paragraph_text(para, "")
            if noi_dung:
                paragraphs_to_insert.append((para, noi_dung))
        
        # Xóa các dòng gạch ngang placeholder _____
        elif text.strip().startswith("_____") and len(text.strip()) > 10:
            _replace_paragraph_text(para, "")
        
        # Kết luận cuộc họp — mục 4.x
        elif text.strip().startswith("4.") and "[Tên mục]:" in text:
            ket_luan_list = session.get("ket_luan", [])
            # Lấy index: "4.1." -> 0, "4.2." -> 1, "4.3." -> 2
            try:
                idx = int(text.strip()[2]) - 1  # "4.1" -> idx=0
            except (IndexError, ValueError):
                idx = -1
            
            if 0 <= idx < len(ket_luan_list):
                kl = ket_luan_list[idx]
                ten_muc = kl.get("ten_muc", "")
                noi_dung_kl = kl.get("noi_dung", "")
                prefix = text.strip()[:4]  # "4.1." hoặc "4.2." hoặc "4.3."
                _replace_paragraph_text(para, f"{prefix} {ten_muc}:")
                # Ghi nhớ để chèn nội dung sau
                if noi_dung_kl:
                    paragraphs_to_insert.append((para, noi_dung_kl))
            elif idx >= len(ket_luan_list):
                # Mục trong template nhưng không có dữ liệu → xóa trắng
                _replace_paragraph_text(para, "")
        
        # Kết thúc cuộc họp
        elif "Cuộc họp kết thúc lúc" in text:
            gio_kt = session.get("gio_ket_thuc", "____")
            ngay_kt = session.get("ngay_ket_thuc", "")
            _replace_paragraph_text(
                para,
                f"Cuộc họp kết thúc lúc {gio_kt}, ngày {ngay_kt or ngay} tháng {thang} năm {nam}."
            )
    
    # Chèn nội dung kết luận & nội dung (phải làm sau khi iterate xong để tránh lỗi)
    for anchor_para, kl_text in reversed(paragraphs_to_insert):
        _insert_paragraph_after(anchor_para, kl_text, doc)
    
    # ========== 3. Fill bảng thành viên (Table 1) ==========
    member_table = doc.tables[1]
    thanh_vien = session.get("thanh_vien", [])
    
    # Table 1 có header row (row 0) + 10 empty rows (row 1-10)
    for i, tv in enumerate(thanh_vien):
        row_idx = i + 1  # Row 0 là header
        
        if row_idx < len(member_table.rows):
            row = member_table.rows[row_idx]
        else:
            row = member_table.add_row()
        
        row.cells[0].text = str(tv.get("stt", i + 1))
        row.cells[1].text = tv.get("ho_ten", "")
        row.cells[2].text = tv.get("chuc_vu", "")
        row.cells[3].text = tv.get("ghi_chu", "")
        
        # Áp dụng font size
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(12)
    
    # Xóa các row trống còn lại
    rows_to_keep = len(thanh_vien) + 1  # +1 cho header
    while len(member_table.rows) > rows_to_keep:
        last_row = member_table.rows[-1]
        tbl = member_table._tbl
        tbl.remove(last_row._tr)
    
    # Xóa bỏ tất cả paragraph trống liên tiếp (dọn dẹp khoảng trắng thừa)
    _remove_consecutive_empty_paragraphs(doc)
    
    return doc


# ==================== Helper Functions ====================

def _replace_paragraph_text(paragraph, new_text: str):
    """
    Thay thế toàn bộ text trong paragraph, giữ nguyên formatting của run đầu tiên.
    """
    if not paragraph.runs:
        paragraph.text = new_text
        return
    
    # Giữ format của run đầu tiên, xóa tất cả run khác
    first_run = paragraph.runs[0]
    for run in paragraph.runs:
        run.text = ""
    first_run.text = new_text


def _replace_in_cell_paragraphs(cell, old_text: str, new_text: str):
    """
    Tìm và thay thế text trong tất cả paragraph của một cell.
    """
    for para in cell.paragraphs:
        if old_text in para.text:
            full_text = para.text.replace(old_text, new_text)
            _replace_paragraph_text(para, full_text)


def _insert_paragraph_after(paragraph, text: str, doc):
    """
    Chèn một hoặc nhiều paragraph mới ngay sau paragraph hiện tại.
    Hỗ trợ text có ký tự xuống dòng (\\n).
    Loại bỏ ký tự markdown và dòng trống.
    """
    if not text:
        return
    
    # Tách dòng, bỏ dòng trống
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    current_anchor = paragraph
    
    for line in lines:
        # Bỏ các ký tự markdown như **, *, #
        line = line.replace('**', '').replace('*', '').replace('#', '').strip()
        if not line:
            continue
        
        new_para = doc.add_paragraph()
        new_para.text = line
        
        # Di chuyển paragraph mới vào sau anchor
        current_anchor._element.addnext(new_para._element)
        
        # Copy style từ anchor
        new_para.style = paragraph.style
        
        # Canh lề: thụt đầu dòng ~1cm, không thụt trái
        pf = new_para.paragraph_format
        pf.first_line_indent = Inches(0.4)
        pf.left_indent = Inches(0)
        # Giảm khoảng cách trước/sau dòng để tránh tạo khoảng trắng lớn
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        
        # Set font size
        for run in new_para.runs:
            run.font.size = Pt(12)
        
        current_anchor = new_para


def _append_doc_with_page_break(main_doc: Document, append_doc: Document):
    """
    Ghép nội dung của append_doc vào cuối main_doc, có chèn Page Break ở giữa.
    Bao gồm cả paragraphs và tables.
    """
    # Chèn page break vào cuối main_doc
    main_doc.add_page_break()
    
    # Copy tất cả element (paragraph + table) từ append_doc sang main_doc
    for element in append_doc.element.body:
        tag = element.tag
        # Bỏ qua sectPr (page settings) để tránh conflict
        if tag.endswith('sectPr'):
            continue
        main_doc.element.body.append(copy.deepcopy(element))


def _remove_consecutive_empty_paragraphs(doc: Document):
    """
    Xóa bỏ các paragraph trống liên tiếp (giữ lại tối đa 1 paragraph trống).
    Giúp dọn dẹp khoảng trắng thừa trong file DOCX.
    """
    body = doc.element.body
    prev_was_empty = False
    elements_to_remove = []
    
    for child in body:
        if child.tag.endswith('}p'):  # Là paragraph
            # Kiểm tra paragraph có text hay không
            text_content = ''.join(node.text or '' for node in child.iter() if node.text)
            is_empty = not text_content.strip()
            
            if is_empty and prev_was_empty:
                elements_to_remove.append(child)
            
            prev_was_empty = is_empty
        else:
            prev_was_empty = False
    
    for elem in elements_to_remove:
        body.remove(elem)
